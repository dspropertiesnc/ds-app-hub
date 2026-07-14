import os, json, base64, io, uuid, tempfile
try:
    import pillow_heif; pillow_heif.register_heif_opener()
except Exception:
    pass
from flask import Blueprint, request, jsonify, render_template, send_file, abort
from werkzeug.utils import secure_filename

bp = Blueprint("contracts", __name__, url_prefix="/contracts", template_folder="templates")

HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.normpath(os.path.join(HERE, "..", "..", "static", "logo.png"))
JOBS = os.environ.get("CONTRACT_TMP", os.path.join(tempfile.gettempdir(), "contract_jobs"))
os.makedirs(JOBS, exist_ok=True)
MODEL = os.getenv("CONTRACT_MODEL", os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"))

META = {"name": "Contract Term Extractor", "desc": "Upload a lease or management agreement; pull out fees, dates, and filled-in terms.",
        "url": "/contracts/", "group": "Leasing & Contracts", "icon": "\U0001F4C4", "ready": True}

SYSTEM = """You extract key terms from a property management agreement (PMA) or residential lease for Doss & Spaulding Properties.
The output is a CONCISE 1-2 page reference sheet an administrator uses to input data into their system.
Be brief: short labels, short values (numbers, %, $, dates, Yes/No). NO sentences, NO explanations, NO legal prose.
Prioritize: all dates & terms, ALL fees and dollar/percent amounts, and values filled into blank fields.

Return STRICT JSON only:
{
  "doc_type": "Property Management Agreement | Residential Lease | Other",
  "property_address": "",
  "sections": [ {"name": "", "rows": [ {"label": "", "value": ""} ] } ]
}

Include ONLY rows whose value is actually present in the document. Omit anything not stated. Keep values terse and exactly as written (e.g. "10.0% of gross rent", "$250", "30 days", "04/10/2026", "Yes", "No").

For a PROPERTY MANAGEMENT AGREEMENT, use these sections and label ideas (skip any not present):
- Owner Information: Name, Email, Phone, Mailing Address, Principal Contacts
- Property Details: Address, County, Legal Description, Water/Sewage
- Agreement Terms: Effective Date, Initial Term, Renewal Term, Termination Notice
- Fee Schedule: Management Fee, Leasing/Placement, Onboarding/Setup, Lease Renewal, Project/Maintenance Oversight, Hourly (out-of-scope), Utility Mgmt Charge, Co-op Agent Cap, Self-Show Opt-Out, Sale-to-Tenant
- Financial Setup: Reserve Fund, Max Repair (no approval), Max Lease Term, Late Fees To, Trust Interest To, Interest Rate, Owner Payment Due
- Insurance: Company, Agent, Contact, Min Liability
- Policies: Pets Allowed, Smoking, HOA, Vacancy Utilities, Self-Showings

For a RESIDENTIAL LEASE, use these sections (skip any not present):
- Tenant Information: Tenant(s), Email, Phone
- Property: Address, Unit
- Lease Terms: Start Date, End Date, Term Length, Renewal, Move-in Date, Notice to Vacate
- Rent & Fees: Monthly Rent, Due Date, Late Fee, NSF Fee, Pet Fee/Rent, Other Fees
- Deposits: Security Deposit, Pet Deposit, Other
- Utilities & Policies: Utilities Paid By, Pets, Smoking, Occupancy Limit, Parking

Output ONLY the JSON object, no prose, no code fences."""

def _pdf_text(data):
    try:
        from pypdf import PdfReader
        import io as _io
        r = PdfReader(_io.BytesIO(data))
        return "\n".join((pg.extract_text() or "") for pg in r.pages)
    except Exception:
        return ""

def _docx_text(path):
    from docx import Document
    d = Document(path); out = []
    for p in d.paragraphs:
        if p.text.strip(): out.append(p.text)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells): out.append(" | ".join(cells))
    return "\n".join(out)[:120000]

def _img_block(fileobj):
    from PIL import Image, ImageOps
    im = Image.open(fileobj); im = ImageOps.exif_transpose(im); im.thumbnail((1600, 1600))
    if im.mode in ("RGBA", "P"): im = im.convert("RGB")
    buf = io.BytesIO(); im.save(buf, "JPEG", quality=80)
    return {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg",
            "data": base64.b64encode(buf.getvalue()).decode()}}

def _parse_json(raw):
    import json as _j
    raw = raw.strip()
    start = raw.find("{")
    if start > 0: raw = raw[start:]
    try:
        return _j.loads(raw)
    except Exception:
        pass
    # attempt to close a truncated object: cut at last complete item, balance braces/brackets
    cut = max(raw.rfind("}"), raw.rfind("]"))
    frag = raw[:cut+1] if cut > 0 else raw
    for _ in range(6):
        try:
            return _j.loads(frag)
        except Exception:
            # drop trailing incomplete fragment after the last comma, then re-balance
            c = frag.rfind(",")
            frag = frag[:c] if c > 0 else frag
            opens = frag.count("{") - frag.count("}")
            openb = frag.count("[") - frag.count("]")
            frag = frag + "]" * max(0, openb) + "}" * max(0, opens)
    return _j.loads(raw)  # re-raise original if unrecoverable

@bp.route("/")
def index():
    return render_template("contracts.html")

@bp.route("/extract", methods=["POST"])
def extract():
    if not os.getenv("ANTHROPIC_API_KEY"):
        return jsonify({"error": "Server has no ANTHROPIC_API_KEY set."}), 500
    files = [f for f in request.files.getlist("files") if f and f.filename]
    if not files:
        return jsonify({"error": "Please upload a contract file."}), 400
    content = []
    for f in files:
        name = f.filename.lower()
        try:
            if name.endswith(".pdf"):
                data = f.read()
                txt = _pdf_text(data)
                if len(txt.strip()) >= 400:  # digital PDF with a real text layer -> fast text path
                    content.append({"type": "text", "text": "[PDF text: %s]\n%s" % (f.filename, txt[:180000])})
                else:  # scanned/image PDF -> let Claude read it visually
                    content.append({"type": "document", "source": {"type": "base64",
                        "media_type": "application/pdf", "data": base64.b64encode(data).decode()}})
            elif name.endswith(".docx"):
                tmp = os.path.join(JOBS, "u_" + uuid.uuid4().hex[:8] + ".docx"); f.save(tmp)
                txt = _docx_text(tmp)
                try: os.remove(tmp)
                except OSError: pass
                content.append({"type": "text", "text": "[Word document contents]\n" + txt})
            else:  # image / scan / photo
                content.append(_img_block(f.stream))
        except Exception as e:
            return jsonify({"error": f"Could not read {f.filename}: {e}"}), 400
    content.append({"type": "text", "text": "Extract the terms from the contract above as specified."})
    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
        msg = client.messages.create(model=MODEL, max_tokens=8000, system=SYSTEM,
                                      messages=[{"role": "user", "content": content}])
        raw = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()
        if raw.startswith("```"):
            raw = raw.strip("`").split("\n", 1)[1].rsplit("```", 1)[0]
        data = _parse_json(raw)
    except Exception as e:
        return jsonify({"error": f"Extraction failed: {e}"}), 502
    job = uuid.uuid4().hex[:12]
    with open(os.path.join(JOBS, job + ".json"), "w") as fh:
        json.dump(data, fh)
    data["job"] = job
    return jsonify(data)

import re as _re
def _summary_filename(data):
    dt = (data.get("doc_type") or "").lower()
    if "lease" in dt:
        prefix = "Lease"
    elif "management" in dt or "pma" in dt:
        prefix = "PMA"
    else:
        prefix = ""
    addr = (data.get("property_address") or data.get("property") or "").strip()
    # ignore non-address placeholders
    if addr and (len(addr) > 90 or addr.lower().startswith("see ")):
        addr = ""
    base = f"{prefix} Contract Summary".strip()
    name = base + (f" - {addr}" if addr else "")
    name = _re.sub(r'[\\/:*?"<>|\r\n]+', " ", name)   # strip filename-illegal chars
    name = _re.sub(r"\s+", " ", name).strip()
    return name[:120] + ".pdf"

@bp.route("/download/<job>")
def download(job):
    job = secure_filename(job)
    jp = os.path.join(JOBS, job + ".json")
    if not os.path.exists(jp): abort(404)
    data = json.load(open(jp))
    pdf_path = os.path.join(JOBS, job + ".pdf")
    _build_pdf(data, pdf_path)
    return send_file(pdf_path, as_attachment=True, download_name=_summary_filename(data))

def _build_pdf(data, out_path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, Image as RLImage, KeepTogether)
    ACCENT = colors.HexColor("#a78147"); SUB = colors.HexColor("#efe7d9")
    DARK = colors.HexColor("#5c4522"); LINE = colors.HexColor("#e6e0d6")
    dt = (data.get("doc_type") or "").lower()
    kind = "PMA Summary" if ("management" in dt or "pma" in dt) else ("Lease Summary" if "lease" in dt else "Contract Summary")
    addr = (data.get("property_address") or data.get("property") or "").strip()

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Title"], fontSize=13, textColor=ACCENT, spaceAfter=0, alignment=0, leading=15)
    sub = ParagraphStyle("sub", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#666666"), spaceAfter=2)
    band = ParagraphStyle("band", parent=styles["Normal"], fontSize=9.5, textColor=colors.white, fontName="Helvetica-Bold")
    lbl = ParagraphStyle("lbl", parent=styles["Normal"], fontSize=8.4, textColor=DARK, fontName="Helvetica-Bold", leading=9.8)
    val = ParagraphStyle("val", parent=styles["Normal"], fontSize=8.4, textColor=colors.HexColor("#1a1a1a"), leading=9.8)

    doc = SimpleDocTemplate(out_path, pagesize=letter, topMargin=0.4*inch, bottomMargin=0.4*inch,
                            leftMargin=0.6*inch, rightMargin=0.6*inch, title=kind)
    W = 7.3*inch
    el = []
    logo = ""
    if os.path.exists(LOGO):
        try: logo = RLImage(LOGO, width=1.5*inch, height=1.5*inch*303/1009)
        except Exception: logo = ""
    head = Table([[Paragraph(kind + ((" &mdash; " + addr) if addr else ""), h1), logo]],
                 colWidths=[W-1.7*inch, 1.7*inch])
    head.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"MIDDLE"), ("ALIGN",(1,0),(1,0),"RIGHT"),
                              ("LEFTPADDING",(0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0)]))
    el.append(head)
    el.append(Paragraph("Doss &amp; Spaulding Properties &nbsp;|&nbsp; auto-extracted reference &mdash; verify against the signed agreement", sub))
    el.append(Spacer(1, 5))

    def band_row(txt):
        t = Table([[Paragraph(txt.upper(), band)]], colWidths=[W])
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),ACCENT), ("LEFTPADDING",(0,0),(-1,-1),6),
                               ("TOPPADDING",(0,0),(-1,-1),1.3), ("BOTTOMPADDING",(0,0),(-1,-1),1.3)]))
        return t

    def rows_table(rows):
        data_rows = []
        for r in rows:
            v = r.get("value")
            if v is None or str(v).strip() == "": continue
            data_rows.append([Paragraph(str(r.get("label") or ""), lbl), Paragraph(str(v), val)])
        if not data_rows: return None
        t = Table(data_rows, colWidths=[2.15*inch, W-2.15*inch])
        t.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP"),
                               ("LINEBELOW",(0,0),(-1,-2),0.3,LINE),
                               ("LEFTPADDING",(0,0),(-1,-1),6), ("RIGHTPADDING",(0,0),(-1,-1),6),
                               ("TOPPADDING",(0,0),(-1,-1),1.3), ("BOTTOMPADDING",(0,0),(-1,-1),1.3)]))
        return t

    for sec in data.get("sections", []):
        rt = rows_table(sec.get("rows", []))
        if rt is None: continue
        el.append(Spacer(1, 3))
        el.append(KeepTogether([band_row(sec.get("name","")), Spacer(1,1), rt]))
    doc.build(el)
    return out_path
