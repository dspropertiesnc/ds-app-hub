"""Builds the styled unit-turn punchlist .docx (Doss & Spaulding house style).
Photos are grouped at the end of each top-level section. No materials list."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="a78147"; SUBBAR="efe7d9"; H2TEXT="6b4f22"; NOTE_GRAY=RGBColor(0x55,0x55,0x55); SUB_GRAY=RGBColor(0x66,0x66,0x66)

def _shade(p,hx):
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hx); pPr.append(sh)
def _sp(p,b,a):
    pf=p.paragraph_format; pf.space_before=Pt(b); pf.space_after=Pt(a)

def _title(doc,t,s,logo_path=None):
    if logo_path and os.path.exists(logo_path):
        tbl=doc.add_table(rows=1,cols=2); tbl.autofit=False
        tbl.columns[0].width=Inches(4.8); tbl.columns[1].width=Inches(2.1)
        left=tbl.rows[0].cells[0]; right=tbl.rows[0].cells[1]
        left.width=Inches(4.8); right.width=Inches(2.1)
        lp=left.paragraphs[0]; _sp(lp,0,0)
        r=lp.add_run(t); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string(ACCENT)
        sp2=left.add_paragraph(); _sp(sp2,2,0)
        r2=sp2.add_run(s); r2.font.size=Pt(10.5); r2.font.color.rgb=SUB_GRAY; r2.italic=True
        rp=right.paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _sp(rp,0,0)
        try: rp.add_run().add_picture(logo_path,width=Inches(1.6))
        except Exception as e: print("logo skip",e)
        doc.add_paragraph(); _sp(doc.paragraphs[-1],0,6)
    else:
        p=doc.add_paragraph(); _sp(p,0,0)
        r=p.add_run(t); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string(ACCENT)
        p2=doc.add_paragraph(); _sp(p2,2,14)
        r2=p2.add_run(s); r2.font.size=Pt(10.5); r2.font.color.rgb=SUB_GRAY; r2.italic=True

def _h1(doc,t):
    p=doc.add_paragraph(); _sp(p,14,6); _shade(p,ACCENT)
    r=p.add_run("  "+t.upper()); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0xff,0xff,0xff)

def _h2(doc,t):
    p=doc.add_paragraph(); _sp(p,10,3); _shade(p,SUBBAR)
    r=p.add_run("  "+t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(H2TEXT)


def _access(doc, text):
    p = doc.add_paragraph(); _sp(p, 10, 2); _shade(p, ACCENT)
    r = p.add_run("  ACCESS INFORMATION"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xff,0xff,0xff)
    for line in [ln for ln in text.splitlines() if ln.strip()]:
        bp = doc.add_paragraph(); _sp(bp, 1, 1); bp.paragraph_format.left_indent = Inches(0.1)
        rr = bp.add_run(line.strip()); rr.font.size = Pt(10.5)

def _item(doc,task,note=None,done=False):
    p=doc.add_paragraph(); _sp(p,3,1); p.paragraph_format.left_indent=Inches(0.1)
    box=p.add_run("☑ " if done else "☐ "); box.font.name="Segoe UI Symbol"; box.font.size=Pt(12)
    rf=box._element.rPr.rFonts; rf.set(qn('w:ascii'),"Segoe UI Symbol"); rf.set(qn('w:hAnsi'),"Segoe UI Symbol")
    t=p.add_run(task); t.font.size=Pt(10.5)
    if done:
        d=p.add_run("   — completed"); d.italic=True; d.font.size=Pt(9.5); d.font.color.rgb=SUB_GRAY
    if note:
        np=doc.add_paragraph(); _sp(np,0,3); np.paragraph_format.left_indent=Inches(0.42)
        nr=np.add_run("Note: "+note); nr.italic=True; nr.font.size=Pt(9.5); nr.font.color.rgb=NOTE_GRAY

def _photos(doc,paths,cols=2,width=3.15):
    paths=[p for p in paths if os.path.exists(p)]
    if not paths: return
    _h2(doc,"Photos")
    rows=(len(paths)+cols-1)//cols
    tbl=doc.add_table(rows=rows,cols=cols); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,path in enumerate(paths):
        cell=tbl.rows[i//cols].cells[i%cols]
        cp=cell.paragraphs[0]; cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        try: cp.add_run().add_picture(path,width=Inches(width))
        except Exception as e: print("skip img",path,e)

def build(spec,out_path,photo_dir="",logo_path=None):
    """spec: {title, subtitle, sections:[{name, subsections:[{name, items:[{task,note,done}]}], photos:[filenames]}]}"""
    doc=Document()
    for s in doc.sections:
        s.top_margin=Inches(0.7); s.bottom_margin=Inches(0.7); s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)
    _title(doc,spec["title"],spec.get("subtitle",""),logo_path)
    if spec.get("access"):
        _access(doc, spec["access"])
    for sec in spec["sections"]:
        _h1(doc,sec["name"])
        for sub in sec.get("subsections",[]):
            _h2(doc,sub["name"])
            for it in sub.get("items",[]):
                _item(doc,it["task"],it.get("note"),it.get("done",False))
        photos=[p if os.path.isabs(p) else os.path.join(photo_dir,p) for p in sec.get("photos",[])]
        _photos(doc,photos)
    doc.save(out_path)
    return out_path
